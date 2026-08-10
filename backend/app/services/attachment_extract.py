"""Extract plain text from uploaded documents (PDF / DOCX / Markdown / text / images).

Extraction is text-only and lossy by design: the output feeds an LLM that
drafts structured entries, staged for user review before any merge.

Image-based documents (certificates, scans, photo exports) carry their real
content as pixels, not a text layer — a certificate PDF often extracts to just
the recipient's name. For those, a vision-model transcription pass turns the
rendered pages into text so every downstream consumer (KB ingest, chat
attachments, entity documents) sees the actual content. The transcription
degrades to the raw text layer when the provider is unavailable.

PDF text comes from pdfplumber (MIT, already the parse-certification engine),
PDF rasterization from pypdfium2 (BSD-3-Clause/Apache-2.0), and standalone
images from Pillow (MIT-CMU) — pypdfium2 is PDF-only. PyMuPDF is deliberately
NOT used anywhere: it is AGPL-3.0 and would relicense the whole project.
"""

import io
import logging
from pathlib import Path

logger = logging.getLogger(__name__)

_TEXT_SUFFIXES = {".md", ".markdown", ".txt", ".tex"}
# A .tex is plain text; some browsers send x-tex mimes, others send an empty
# mime (routed by the .tex suffix above).
_TEXT_MIMES = {"text/plain", "text/markdown", "application/x-tex", "text/x-tex"}

_IMAGE_SUFFIXES = {".png", ".jpg", ".jpeg", ".webp"}
_IMAGE_MIMES = {"image/png", "image/jpeg", "image/webp"}

MAX_CHARS = 200_000

# A PDF whose text layer is this sparse AND that embeds images is treated as
# image-based (scanned page, generated certificate) and sent to vision.
VISION_TEXT_THRESHOLD = 150
VISION_MAX_PAGES = 5
VISION_MARKER = "[Transcribed from document images]"

_TRANSCRIBE_PROMPT = (
    "Transcribe ALL text visible in the attached document page image(s), faithfully and "
    "completely: titles, headings, names, organizations, dates, credential IDs, URLs, and "
    "body text. Preserve the reading order. Output plain text only — no commentary, no "
    "markdown fences, no descriptions of decorative graphics."
)


def _pdf_text(data: bytes) -> str:
    import pdfplumber

    with pdfplumber.open(io.BytesIO(data)) as doc:
        return "\n\n".join(page.extract_text() or "" for page in doc.pages)


def _pdf_has_images(data: bytes) -> bool:
    import pdfplumber

    with pdfplumber.open(io.BytesIO(data)) as doc:
        return any(page.images for page in doc.pages)


def _pdf_page_pngs(data: bytes, scale: float) -> list[bytes]:
    """Render up to VISION_MAX_PAGES PDF pages to PNG bytes."""
    import pypdfium2 as pdfium

    out: list[bytes] = []
    pdf = pdfium.PdfDocument(data)
    try:
        for i in range(min(len(pdf), VISION_MAX_PAGES)):
            buf = io.BytesIO()
            pdf[i].render(scale=scale).to_pil().save(buf, format="PNG")
            out.append(buf.getvalue())
    finally:
        pdf.close()
    return out


def _image_png(data: bytes) -> list[bytes]:
    """Normalize a standalone image to a single PNG.

    pypdfium2 is PDF-only, so images take the Pillow path. Opening here also
    validates the bytes — the caller turns a failure into a 4xx.
    """
    from PIL import Image

    with Image.open(io.BytesIO(data)) as img:
        img.load()
        # PNG cannot hold CMYK and friends; normalize anything exotic.
        if img.mode not in ("1", "L", "LA", "P", "RGB", "RGBA"):
            img = img.convert("RGB")
        buf = io.BytesIO()
        img.save(buf, format="PNG")
    return [buf.getvalue()]


def _transcribe_images(images: list[bytes], filename: str) -> str:
    """OCR-via-vision-model; returns "" on any failure so callers can degrade."""
    from app.services import llm, model_settings

    try:
        text = llm.call_openai(
            prompt=_TRANSCRIBE_PROMPT,
            model=model_settings.get_fast_model(),
            response_format="text",
            images=images,
            trace_name="doc-vision-transcribe",
        )
    except Exception as exc:  # noqa: BLE001 — degrade to the raw text layer
        logger.warning("Vision transcription failed for %r: %s", filename, exc)
        return ""
    return (text or "").strip()


def _docx_text(data: bytes) -> str:
    import docx

    document = docx.Document(io.BytesIO(data))
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(parts)


def extract_text(filename: str, mime: str | None, data: bytes) -> str:
    suffix = Path(filename or "").suffix.lower()
    mime = (mime or "").lower().split(";")[0].strip()

    if suffix == ".pdf" or mime == "application/pdf":
        text = _pdf_text(data)
        if len(text.strip()) < VISION_TEXT_THRESHOLD and _pdf_has_images(data):
            transcript = _transcribe_images(_pdf_page_pngs(data, scale=2.0), filename)
            if transcript:
                text = f"{text.strip()}\n\n{VISION_MARKER}\n{transcript}"
    elif suffix == ".docx" or mime == (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    ):
        text = _docx_text(data)
    elif suffix in _IMAGE_SUFFIXES or mime in _IMAGE_MIMES:
        try:
            pages = _image_png(data)
        except Exception as exc:  # corrupt/undecodable image bytes
            raise ValueError(f"Attachment {filename!r} is not a readable image: {exc}") from exc
        transcript = _transcribe_images(pages, filename)
        text = f"{VISION_MARKER}\n{transcript}" if transcript else ""
    elif suffix in _TEXT_SUFFIXES or mime in _TEXT_MIMES:
        text = data.decode("utf-8", errors="replace")
    else:
        raise ValueError(
            f"Unsupported attachment type: {filename!r} ({mime or 'unknown mime'}). "
            "Supported: .pdf, .docx, .md, .txt, .tex, .png, .jpg, .webp"
        )

    text = text.strip()
    if not text:
        raise ValueError(f"Attachment {filename!r} contained no extractable text")
    return text[:MAX_CHARS]
